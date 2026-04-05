#!/usr/bin/env python3
"""
create-reference-docx.py
Generates a professional reference.docx for Pandoc document generation.

The reference.docx acts as a Word style template. Pandoc maps Markdown
elements to the named styles defined here:
  # Heading 1  → "Heading 1"
  ## Heading 2 → "Heading 2"
  Body text    → "Normal"
  Code block   → "Source Code" / "Verbatim Char"
  Table        → "Table" styles

Usage:
  python3 scripts/create-reference-docx.py
  Outputs: templates/reference.docx
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# Colour palette — professional corporate navy/blue
# ---------------------------------------------------------------------------
NAVY       = RGBColor(0x1F, 0x38, 0x64)   # Heading 1 — dark navy
STEEL      = RGBColor(0x2E, 0x5D, 0x8A)   # Heading 2 — steel blue
ACCENT     = RGBColor(0x44, 0x72, 0xC4)   # Heading 3 — accent blue
DARK_GREY  = RGBColor(0x26, 0x26, 0x26)   # Body text
MID_GREY   = RGBColor(0x59, 0x59, 0x59)   # Subtitle / caption
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)   # Table alt-row shading
CODE_BG    = RGBColor(0xF5, 0xF5, 0xF5)   # Code block background
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# Font names
# ---------------------------------------------------------------------------
FONT_HEADING = "Calibri Light"
FONT_BODY    = "Calibri"
FONT_CODE    = "Consolas"

# ---------------------------------------------------------------------------
# Helper — set paragraph shading (background colour)
# ---------------------------------------------------------------------------
def set_paragraph_shading(paragraph, fill_color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color_hex)
    pPr.append(shd)

# ---------------------------------------------------------------------------
# Helper — set run shading
# ---------------------------------------------------------------------------
def set_run_shading(run, fill_color_hex):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color_hex)
    rPr.append(shd)

# ---------------------------------------------------------------------------
# Helper — add left border to paragraph (used on Heading 1)
# ---------------------------------------------------------------------------
def add_left_border(paragraph, color_hex="1F3864", space="4", sz="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

# ---------------------------------------------------------------------------
# Helper — set exact line spacing in twips
# ---------------------------------------------------------------------------
def set_line_spacing(paragraph_format, rule, value_pt):
    paragraph_format.line_spacing_rule = rule
    paragraph_format.line_spacing = Pt(value_pt)

# ---------------------------------------------------------------------------
# Helper — set table cell background
# ---------------------------------------------------------------------------
def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

# ---------------------------------------------------------------------------
# Build the reference document
# ---------------------------------------------------------------------------
def build_reference_docx(output_path):
    doc = Document()

    # --- Page layout ---
    section = doc.sections[0]
    section.page_width  = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.00)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)
    section.different_first_page_header_footer = True

    # --- Default font ---
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK_GREY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # --- Heading 1 ---
    h1 = styles["Heading 1"]
    h1.font.name = FONT_HEADING
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False
    # Left bar accent
    pPr = h1._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "1F3864")
    pBdr.append(left)
    pPr.append(pBdr)

    # --- Heading 2 ---
    h2 = styles["Heading 2"]
    h2.font.name = FONT_HEADING
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = STEEL
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # --- Heading 3 ---
    h3 = styles["Heading 3"]
    h3.font.name = FONT_BODY
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = ACCENT
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(3)
    h3.paragraph_format.keep_with_next = True

    # --- Heading 4 ---
    h4 = styles["Heading 4"]
    h4.font.name = FONT_BODY
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = DARK_GREY
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after  = Pt(2)
    h4.paragraph_format.keep_with_next = True

    # --- Title ---
    try:
        title_style = styles["Title"]
    except KeyError:
        title_style = styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = FONT_HEADING
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = NAVY
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after  = Pt(8)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Subtitle ---
    try:
        sub_style = styles["Subtitle"]
    except KeyError:
        sub_style = styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    sub_style.font.name = FONT_BODY
    sub_style.font.size = Pt(13)
    sub_style.font.italic = True
    sub_style.font.color.rgb = MID_GREY
    sub_style.paragraph_format.space_before = Pt(0)
    sub_style.paragraph_format.space_after  = Pt(24)

    # --- Body Text ---
    try:
        body_style = styles["Body Text"]
    except KeyError:
        body_style = styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
    body_style.base_style = styles["Normal"]
    body_style.font.name  = FONT_BODY
    body_style.font.size  = Pt(11)
    body_style.paragraph_format.space_after  = Pt(8)
    body_style.paragraph_format.space_before = Pt(0)

    # --- First Paragraph ---
    try:
        fp_style = styles["First Paragraph"]
    except KeyError:
        fp_style = styles.add_style("First Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    fp_style.base_style = styles["Normal"]
    fp_style.paragraph_format.first_line_indent = Pt(0)

    # --- Verbatim / Code ---
    for sname in ["Verbatim Char", "Verbatim", "Source Code"]:
        try:
            code_s = styles[sname]
        except KeyError:
            if sname in ["Verbatim Char"]:
                code_s = styles.add_style(sname, WD_STYLE_TYPE.CHARACTER)
            else:
                code_s = styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
        code_s.font.name = FONT_CODE
        code_s.font.size = Pt(9.5)
        code_s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        if hasattr(code_s, "paragraph_format"):
            code_s.paragraph_format.space_before = Pt(4)
            code_s.paragraph_format.space_after  = Pt(4)

    # --- Block Text (blockquotes) ---
    try:
        bq_style = styles["Block Text"]
    except KeyError:
        bq_style = styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
    bq_style.font.name   = FONT_BODY
    bq_style.font.size   = Pt(11)
    bq_style.font.italic = True
    bq_style.font.color.rgb = MID_GREY
    bq_style.paragraph_format.left_indent  = Cm(1.0)
    bq_style.paragraph_format.right_indent = Cm(1.0)
    bq_style.paragraph_format.space_before = Pt(6)
    bq_style.paragraph_format.space_after  = Pt(6)

    # --- Caption ---
    try:
        cap_style = styles["Caption"]
    except KeyError:
        cap_style = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    cap_style.font.name   = FONT_BODY
    cap_style.font.size   = Pt(9)
    cap_style.font.italic = True
    cap_style.font.color.rgb = MID_GREY
    cap_style.paragraph_format.space_before = Pt(2)
    cap_style.paragraph_format.space_after  = Pt(8)
    cap_style.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER

    # --- Table styles ---
    # Pandoc uses "Table" style; we set the default table style
    try:
        tbl_style = styles["Table Grid"]
    except KeyError:
        pass

    # --- Footer: Page X of Y ---
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style = doc.styles["Normal"]
    footer_run = footer_para.add_run()
    footer_run.font.name = FONT_BODY
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MID_GREY

    # Add "Page " field
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run_xml = footer_run._r
    run_xml.append(fldChar_begin)
    run_xml.append(instrText)
    run_xml.append(fldChar_sep)
    run_xml.append(fldChar_end)

    # " of " between page and numpages
    mid_run = footer_para.add_run(" of ")
    mid_run.font.name = FONT_BODY
    mid_run.font.size = Pt(9)
    mid_run.font.color.rgb = MID_GREY

    # NUMPAGES field
    np_run = footer_para.add_run()
    np_run.font.name = FONT_BODY
    np_run.font.size = Pt(9)
    np_run.font.color.rgb = MID_GREY
    np_xml = np_run._r
    np_begin = OxmlElement("w:fldChar")
    np_begin.set(qn("w:fldCharType"), "begin")
    np_instr = OxmlElement("w:instrText")
    np_instr.set(qn("xml:space"), "preserve")
    np_instr.text = " NUMPAGES "
    np_sep = OxmlElement("w:fldChar")
    np_sep.set(qn("w:fldCharType"), "separate")
    np_end = OxmlElement("w:fldChar")
    np_end.set(qn("w:fldCharType"), "end")
    np_xml.append(np_begin)
    np_xml.append(np_instr)
    np_xml.append(np_sep)
    np_xml.append(np_end)

    # --- Header: TITLE field (right-aligned) ---
    header = section.header
    header_para = header.paragraphs[0]
    header_para.clear()
    header_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.style = doc.styles["Normal"]

    # Thin horizontal rule under header via bottom border
    pPr_h = header_para._p.get_or_add_pPr()
    pBdr_h = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr_h.append(bottom)
    pPr_h.append(pBdr_h)

    title_run = header_para.add_run()
    title_run.font.name = FONT_BODY
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = MID_GREY
    # TITLE field
    hdr_xml = title_run._r
    t_begin = OxmlElement("w:fldChar")
    t_begin.set(qn("w:fldCharType"), "begin")
    t_instr = OxmlElement("w:instrText")
    t_instr.set(qn("xml:space"), "preserve")
    t_instr.text = " TITLE "
    t_sep = OxmlElement("w:fldChar")
    t_sep.set(qn("w:fldCharType"), "separate")
    t_end = OxmlElement("w:fldChar")
    t_end.set(qn("w:fldCharType"), "end")
    hdr_xml.append(t_begin)
    hdr_xml.append(t_instr)
    hdr_xml.append(t_sep)
    hdr_xml.append(t_end)

    # --- Write file ---
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[OK] Professional reference.docx written to: {output_path}")


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else "templates/reference.docx"
    build_reference_docx(output)
